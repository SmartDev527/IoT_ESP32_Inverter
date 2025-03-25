import markdown
from docx import Document
from html.parser import HTMLParser
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class MyHTMLParser(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.in_list = False

    def handle_starttag(self, tag, attrs):
        if tag == 'h1':
            self.doc.add_heading('', level=1)
        elif tag == 'h2':
            self.doc.add_heading('', level=2)
        elif tag == 'h3':
            self.doc.add_heading('', level=3)
        elif tag == 'h4':
            self.doc.add_heading('', level=4)
        elif tag == 'ul':
            self.in_list = True
        elif tag == 'li' and self.in_list:
            self.doc.add_paragraph('', style='List Bullet')

    def handle_endtag(self, tag):
        if tag == 'ul':
            self.in_list = False

    def handle_data(self, data):
        data = data.strip()
        if data:
            last_para = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
            if last_para.text == '':
                last_para.text = data
            else:
                self.doc.add_paragraph(data)

def md_to_docx(md_file, docx_file):
    # Read Markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Convert Markdown to HTML
    html = markdown.markdown(md_text, extensions=['extra'])

    # Create Word document
    doc = Document()

    # Parse HTML and populate document
    parser = MyHTMLParser(doc)
    parser.feed(html)

    # Save document
    doc.save(docx_file)

if __name__ == "__main__":
    md_to_docx('ESP32 SIM7600 IoT Device Guide with Python OTA Server.md', 'ESP32 SIM7600 IoT Device Guide with Python OTA Server.docx')